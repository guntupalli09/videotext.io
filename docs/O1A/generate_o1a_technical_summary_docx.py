"""
Generate submission-ready Word document for O-1A Criterion 3 technical summary.
Requires: pip install python-docx (already available in this environment).
Output: RESEARCH_TO_PRODUCT_CRITERION3_TECHNICAL_SUMMARY.docx alongside this script.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_shading(cell, fill_hex: str) -> None:
    """Light gray header row (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_runs(p, parts: list[tuple[str, bool, bool]]) -> None:
    """parts: (text, bold, italic)"""
    for text, bold, italic in parts:
        if not text:
            continue
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name in ("Heading 1", "Heading 2"):
        if name in doc.styles:
            s = doc.styles[name]
            s.font.name = "Calibri"
            s.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            if name == "Heading 1":
                s.font.size = Pt(16)
                s.font.bold = True
            else:
                s.font.size = Pt(13)
                s.font.bold = True


def repo_root_from_script(script_dir: Path) -> Path:
    """docs/O1A -> workspace root."""
    return script_dir.parent.parent


def parse_yaml_frontmatter_simple(md_text: str) -> dict[str, str]:
    """Extract slug and title from leading --- block (minimal parser)."""
    if not md_text.startswith("---"):
        return {}
    end = md_text.find("\n---", 3)
    if end == -1:
        return {}
    block = md_text[3:end]
    out: dict[str, str] = {}
    for line in block.splitlines():
        m = re.match(r"^(slug|title):\s*(.+)\s*$", line)
        if not m:
            continue
        key, val = m.group(1), m.group(2).strip()
        if (val.startswith('"') and val.endswith('"')) or (val.startswith("'") and val.endswith("'")):
            val = val[1:-1]
        out[key] = val
    return out


def collect_videotext_blog_rows(repo_root: Path) -> list[list[str]]:
    """Title, URL, Channel for each content/blog/*.md with slug."""
    blog_dir = repo_root / "content" / "blog"
    rows: list[list[str]] = []
    if not blog_dir.is_dir():
        return rows
    for path in sorted(blog_dir.glob("*.md")):
        try:
            text = path.read_text(encoding="utf-8")
        except OSError:
            continue
        fm = parse_yaml_frontmatter_simple(text)
        slug = fm.get("slug") or path.stem
        title = fm.get("title") or path.stem.replace("-", " ").title()
        url = f"https://blog.videotext.io/{slug}"
        rows.append([title, url, "VideoText.io blog"])
    rows.sort(key=lambda r: r[0].lower())
    return rows


def load_extra_public_writing(script_dir: Path) -> list[list[str]]:
    """
    Optional file: extra_public_writing_urls.txt
    Lines: Title | URL | Platform (platform optional).
    """
    path = script_dir / "extra_public_writing_urls.txt"
    if not path.is_file():
        return []
    rows: list[list[str]] = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        parts = [p.strip() for p in line.split("|")]
        if len(parts) >= 2:
            title, url = parts[0], parts[1]
            platform = parts[2] if len(parts) > 2 else "Syndicated"
            if title and url:
                rows.append([title, url, platform])
        elif len(parts) == 1 and parts[0].startswith("http"):
            rows.append([parts[0], parts[0], "Syndicated"])
    return rows


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(12 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(6)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths_inches: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], "D9D9D9")
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(10)
    if col_widths_inches:
        for row in table.rows:
            for ci, w in enumerate(col_widths_inches):
                if ci < len(row.cells):
                    row.cells[ci].width = Inches(w)


def build_document(script_dir: Path | None = None) -> Document:
    script_dir = script_dir or Path(__file__).resolve().parent
    repo_root = repo_root_from_script(script_dir)

    doc = Document()
    configure_styles(doc)

    # --- Title block ---
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "Technical Summary: ICCS 2026 Research and VideoText.io"
    )
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Calibri"

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run(
        "O-1A Nonimmigrant Classification — Criterion 3: Original Contributions of Major Significance"
    )
    sr.italic = True
    sr.font.size = Pt(12)
    sr.font.name = "Calibri"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "Applicant: Santhosh Guntupalli  |  Product: VideoText.io  |  Memorandum type: Technical exhibit (research ↔ production)"
    )
    mr.font.size = Pt(10)
    mr.font.name = "Calibri"

    doc.add_paragraph()

    # --- Publication & product ---
    add_heading(doc, "Publication and product", level=2)
    p = doc.add_paragraph()
    add_runs(
        p,
        [
            ("Publication (camera-ready): ", True, False),
            (
                "Santhosh Guntupalli, Deterministic Execution Frameworks for Hybrid "
                "Symbolic–Probabilistic Computational Pipelines (ICCS 2026).",
                False,
                True,
            ),
        ],
    )
    p2 = doc.add_paragraph()
    add_runs(
        p2,
        [
            ("Product: ", True, False),
            (
                "VideoText.io — production SaaS for video transcription, subtitles, translation, "
                "caption tooling, batch processing, and related professional media workflows.",
                False,
                False,
            ),
        ],
    )

    add_heading(doc, "Third-party product presence", level=2)
    intro = doc.add_paragraph(
        "The product is publicly listed or featured on the following independent channels "
        "(attach screenshots or stable URLs as exhibits to the petition)."
    )
    intro.paragraph_format.space_after = Pt(6)

    bullets = [
        (
            "IndieHunt",
            "VideoText.io is featured on IndieHunt (independent product discovery).",
        ),
        (
            "Fazier",
            "The product carries the Fazier badge and/or a published listing on the Fazier platform.",
        ),
        (
            "G2",
            "A product profile is published on G2 (https://www.g2.com/) — business software reviews and comparison.",
        ),
        (
            "Capterra",
            "A product profile is published on Capterra (https://www.capterra.com/) — software discovery and reviews.",
        ),
    ]
    for title, body in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        add_runs(bp, [(f"{title}. ", True, False), (body, False, False)])

    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Pt(18)
    nr = note.add_run(
        "Immigration packet: include direct links to each listing and/or dated PDF screenshots "
        "as labeled exhibits (e.g., Exhibit C-3a through C-3d)."
    )
    nr.italic = True
    nr.font.size = Pt(10)

    # --- Current production snapshot ---
    add_heading(doc, "Current production snapshot (field-level significance)", level=2)
    snap_intro = doc.add_paragraph()
    snap_intro.add_run(
        "The following figures are drawn directly from VideoText.io’s internal operations "
        "dashboard (sourced from the canonical Postgres business_users / business_jobs views "
        "described in the analytics-architecture section below) as of 2026-08-08. They are "
        "offered as direct evidence that the applicant’s engineering choices — deterministic, "
        "non-LLM stages for the core pipeline; explicit tier-aware orchestration; result caching — "
        "translate into fast, reliable, well-received service under real usage, not only "
        "theoretical design."
    )

    snap_headers = ["Metric", "Value", "Window"]
    snap_rows = [
        ["Registered users", "489", "cumulative, as of snapshot date"],
        ["New signups", "201 (+67% week-over-week)", "trailing 30 days"],
        ["Jobs completed across all tools", "≈1,480", "trailing 30 days"],
        ["Average job processing time", "24.5 s", "trailing 30 days, completed jobs"],
        ["P95 job processing time", "116.5 s", "trailing 30 days, completed jobs"],
        ["Job failure rate", "5.1%", "trailing 30 days"],
        ["Average customer rating", "4.6 / 5 (38 ratings)", "cumulative"],
        ["Top tool by volume", "Video → Subtitles (153 jobs)", "trailing 30 days"],
    ]
    add_table(doc, snap_headers, snap_rows, col_widths_inches=[2.3, 2.3, 1.9])

    snap_note = doc.add_paragraph()
    snap_note.paragraph_format.space_before = Pt(6)
    snr = snap_note.add_run(
        "These are operational and product-quality metrics (throughput, reliability, customer "
        "satisfaction, and growth), not revenue or subscriber-count figures — counsel may attach a "
        "dated screenshot of this dashboard view as a labeled exhibit, and should refresh the figures "
        "immediately before filing since they change daily."
    )
    snr.italic = True
    snr.font.size = Pt(10)

    # --- Executive summary ---
    add_heading(doc, "Executive summary", level=1)

    es1 = doc.add_paragraph()
    es1.add_run(
        "The ICCS paper contributes a computational framework for hybrid pipelines: a deterministic "
        "execution layer (normalized input, ordered logic, versioned configuration) that owns all state "
        "transitions and primary outputs, while stochastic models are confined to non-authoritative "
        "post-hoc layers that cannot alter deterministic findings. The paper validates this design "
        "with 100% execution determinism and 100% traceability on a structured-text corpus, contrasted "
        "with 0% determinism for pure LLM and schema-constrained LLM baselines, and reports verifiable "
        "execution traces, O(n) rule scaling, and explicit security boundaries (for example, prompt "
        "injection treated as data in the symbolic core)."
    )

    es2 = doc.add_paragraph()
    es2.add_run(
        "VideoText.io is a commercial deployment of the same architectural discipline in a different "
        "domain (professional video and timed text). The product’s backbone is an explicit "
        "upload → validate → queue → worker → download pipeline with tier-aware orchestration. "
        "Deterministic, non-LLM stages (FFmpeg-based media transforms, deterministic transcript "
        "normalization, structured exports keyed off fixed segments, deterministic cache keys) bound "
        "the repeatable, customer-facing artifacts. Optional LLM features (for example, transcript "
        "summary, chapter titling from timestamped segments) are computed only after the primary "
        "timed transcript exists and are stored and delivered as separate fields, so they do not "
        "redefine the core subtitle or transcript payload that customers treat as authoritative for "
        "editing, compliance-style review, or downstream automation. Purely rule-based product logic "
        "(for example, in-app guidance triggers) is implemented without AI, mirroring the paper’s "
        "emphasis on traceable, reproducible decision paths where human operators expect stability. "
        "On raw throughput, the same production pipeline processes long-form video at roughly "
        "47× realtime on a representative nearly-two-hour file — a reproducible, code-instrumented "
        "benchmark (server-side timing logs, not a marketing estimate) that is competitive with or "
        "faster than single-node commercial transcription throughput reported elsewhere in the "
        "industry, evidence that the deterministic engineering discipline below did not come at the "
        "cost of production performance."
    )

    es2b = doc.add_paragraph()
    es2b.add_run(
        "The applicant has independently applied the same discipline — a single deterministic source "
        "of truth, explicit boundaries around any judgment-dependent or probabilistic layer, and "
        "continuous automated reconciliation rather than one-time manual audit — to a second, unrelated "
        "system: VideoText.io’s internal business-analytics platform. That program defines a canonical, "
        "governed metric layer (one definition, one source, one code path per KPI), separates a "
        "deterministic system of record (Postgres/Stripe) from a probabilistic, explicitly-labeled "
        "layer (browser-based session analytics), and requires every reconciliation check to pass with "
        "zero unexplained discrepancies before a metric is allowed to go live. This is not the same "
        "artifact as the ICCS paper, but it is the same underlying engineering judgment restated in a "
        "second domain — evidence of a repeatable methodology rather than a single result."
    )

    es3 = doc.add_paragraph()
    es3.add_run(
        "Neural automatic speech recognition (Whisper-class models) is not the paper’s symbolic rule "
        "engine; spoken audio lacks a closed finite pattern inventory like templated contracts. The "
        "accurate linkage for immigration purposes is therefore not that “the SaaS re-runs the contract "
        "engine,” but rather: (i) the applicant’s peer-reviewed framework establishes execution-invariance "
        "and auditability as first-class engineering constraints for LLM-containing systems; (ii) VideoText.io "
        "applies those constraints where production reality allows (deterministic media and text processing, "
        "explicit pipeline ordering, non-mutating generative add-ons, reproducible caching); (iii) paying "
        "customers demonstrate field-level significance that purely academic artifacts cannot, by showing "
        "the same researcher ships defensible hybrid AI systems under real latency, privacy, billing, and "
        "support obligations."
    )

    # --- Mapping table ---
    add_heading(doc, "Technical mapping: paper concept → VideoText.io implementation", level=1)

    mapping_headers = ["Paper concept", "Role in the framework", "VideoText.io analogue (production)"]
    mapping_rows = [
        [
            "Deterministic execution boundary",
            "All authoritative state transitions and structured outputs live inside a verifiable layer.",
            "FFmpeg pipelines (extract, merge, burn, compress); deterministic transcript post-processing "
            "(transcriptFormatter — no LLM); structured exports built from fixed segment lists.",
        ],
        [
            "Formal execution identity S = (D, R, θ, σ)",
            "Identical execution state implies identical deterministic outputs.",
            "Result cache keyed by user + content hash + tool type + sorted options (computeOptionsHash), "
            "so repeat jobs with identical inputs resolve to the same cached artifact when caching is enabled.",
        ],
        [
            "Stochastic layer as post-hoc explainer",
            "The LLM cannot introduce or change deterministic classifications.",
            "Transcript text and segments are materialized first; optional summary and chapters run afterward "
            "and attach as separate metadata (including deferred summary when Redis is available so the "
            "primary download path stays transcript-first). Chapter generation is constrained to timestamps "
            "present in deterministic segment input.",
        ],
        [
            "Traceability / auditability",
            "Every decision is attributable to logic and input spans.",
            "Timestamped segments, JSON exports, job and request identifiers, operational logging — suitable "
            "for professional workflows that require accountability of outputs.",
        ],
        [
            "Rule-based deterministic control",
            "Pattern logic without stochastic decoding.",
            "Rule-based, deterministic in-app triggers (for example, Tex workflow suggestions: no model inference).",
        ],
        [
            "Domain note in paper",
            "The framework is described as domain-agnostic; the contract study is evaluation, not the only application.",
            "VideoText extends the engineering pattern to media and timed-text pipelines, where reproducibility "
            "applies to encoding, formatting, exports, and orchestration, while ASR remains a known stochastic "
            "component bounded by the pipeline contract above.",
        ],
    ]
    add_table(doc, mapping_headers, mapping_rows, col_widths_inches=[1.35, 1.85, 3.0])

    # --- Second independent application ---
    add_heading(
        doc,
        "A second, independent application of the same engineering discipline",
        level=1,
    )
    sa1 = doc.add_paragraph()
    sa1.add_run(
        "Criterion 3 is strengthened when the same original methodology recurs in an independent "
        "setting, because that shows a transferable engineering judgment rather than a single "
        "publication-shaped result. VideoText.io provides that recurrence internally: separate from "
        "the media-processing pipeline mapped above, the applicant designed and is deploying a "
        "canonical business-analytics architecture governing how every product and revenue metric is "
        "defined, computed, and served."
    )
    sa2 = doc.add_paragraph()
    add_runs(
        sa2,
        [
            ("Same core pattern, different domain. ", True, False),
            (
                "The analytics architecture separates a single deterministic system of record "
                "(Postgres for identity and product usage, Stripe for money) from a probabilistic, "
                "explicitly-labeled layer (browser-based session analytics), and requires every metric "
                "to resolve to exactly one canonical definition, one source, and one code path — the "
                "same structural move as the paper’s separation of a deterministic execution boundary "
                "from a non-authoritative stochastic layer.",
                False,
                False,
            ),
        ],
    )
    sa3 = doc.add_paragraph()
    add_runs(
        sa3,
        [
            ("Independently measured, not just designed. ", True, False),
            (
                "The canonical layer was validated field-by-field against the raw system of record "
                "before any dashboard was allowed to read from it, with every reconciled field passing "
                "with zero unexplained discrepancies — an empirical traceability check in the same "
                "spirit as the paper’s measured 100% traceability result, applied here to a live "
                "commercial system rather than a benchmark corpus.",
                False,
                False,
            ),
        ],
    )
    sa4 = doc.add_paragraph()
    add_runs(
        sa4,
        [
            ("Governed, not one-time. ", True, False),
            (
                "New metrics require a documented definition, a declared source, an owner, and an "
                "automated validation rule before they ship, and any metric that depends on judgment "
                "(for example, model-derived lifetime value or cost-allocation-dependent margin) is "
                "explicitly labeled at reduced confidence rather than presented with false precision — "
                "the same intellectual honesty about the boundary of deterministic guarantees that the "
                "paper itself insists on.",
                False,
                False,
            ),
        ],
    )
    sa5 = doc.add_paragraph()
    sa5.add_run(
        "For adjudication purposes, this second application is offered as corroborating evidence of "
        "originality and consistent engineering judgment, not as a second scholarly contribution. It "
        "shows the applicant applying the peer-reviewed framework’s central idea — a governed "
        "deterministic core with clearly bounded, non-authoritative probabilistic layers — a second "
        "time, on a different system, at his own initiative, in production."
    )

    # --- Criterion 3 ---
    add_heading(doc, "Why this supports “major significance” under Criterion 3", level=1)

    o1 = doc.add_paragraph()
    o1.add_run("1. Originality (research). ").bold = True
    o1.add_run(
        "The ICCS work is not a prompt tweak; it is a reproducibility-first execution model with empirical "
        "measurement of determinism, traceability, variance, and cost, including negative results for "
        "“structured output” LLM baselines that still fail execution invariance."
    )

    o2 = doc.add_paragraph()
    o2.add_run("2. Significance (industry). ").bold = True
    o2.add_run(
        "VideoText.io converts that research mindset into shipping software with subscriptions and usage "
        "metering, proving the applicant can operationalize hybrid AI systems beyond a lab benchmark — a "
        "relatively rare combination in O-1A filings that often show either publications or a product, not "
        "a coherent bridge between the two. That the applicant then independently re-applied the same "
        "governed-determinism methodology to a second, unrelated system (the analytics architecture "
        "described above) — reaching a measured zero-unexplained-discrepancy reconciliation result without "
        "being asked to — indicates a repeatable engineering judgment rather than a one-off result, and the "
        "production pipeline’s own throughput (approximately 47× realtime on long-form video, measured, not "
        "estimated) shows the discipline was applied without sacrificing commercial performance. Independent "
        "distribution and review channels reinforce field recognition: the product is featured on IndieHunt, "
        "carries a Fazier badge or listing, and is published on G2 and Capterra, which are widely used by "
        "buyers to evaluate software — evidence that the contribution sits in the commercial marketplace, "
        "not only in a private deployment."
    )

    o3 = doc.add_paragraph()
    o3.add_run("3. Integrity for adjudication. ").bold = True
    o3.add_run(
        "The strongest argument is architectural and professional continuity (deterministic cores where feasible, "
        "explicit boundaries for generative layers, traceable artifacts, production hardening), plus the paper’s "
        "quantitative guarantees in its evaluated domain. Claiming VideoText reproduces the paper’s 100% "
        "determinism for speech transcription would be inaccurate and should not be asserted; the accurate claim "
        "is controlled hybrid design and commercial proof of impact grounded in the same research-led engineering "
        "philosophy."
    )

    # --- Exhibit checklist ---
    add_heading(doc, "Third-party recognition — exhibit checklist", level=1)
    doc.add_paragraph(
        "Use the following table when assembling Criterion 3 evidence; keep URLs current on the exhibit cover sheet."
    )

    ex_headers = ["Channel", "Claim in this memorandum", "Suggested exhibit"]
    ex_rows = [
        [
            "IndieHunt",
            "Featured listing or feature placement for VideoText.io",
            "Screenshot or permalink to the IndieHunt feature page",
        ],
        [
            "Fazier",
            "Fazier badge and/or published product page",
            "Screenshot of badge on site and/or Fazier product URL",
        ],
        [
            "G2",
            "Published product profile",
            "G2 product page URL plus dated screenshot",
        ],
        [
            "Capterra",
            "Published product profile",
            "Capterra product page URL plus dated screenshot",
        ],
    ]
    add_table(doc, ex_headers, ex_rows, col_widths_inches=[1.1, 2.2, 2.9])

    # --- Appendix: public writing (blogs, Medium, etc.) ---
    add_heading(doc, "Appendix: Public writing supporting VideoText.io", level=1)
    ap1 = doc.add_paragraph()
    ap1.add_run(
        "The following items are public-facing articles used for product education, SEO, and distribution. "
        "They are not peer-reviewed scholarship; they supplement this memorandum by showing sustained, "
        "field-oriented communication around the same product described above. Immigration counsel may "
        "treat this appendix as part of this exhibit or split blog URLs and Medium reprints into separate "
        "labeled exhibits, depending on petition strategy."
    )
    ap2 = doc.add_paragraph()
    ap2.add_run(
        "Recommendation: do not paste full article text into this Word file—keep the packet readable. "
        "Provide titles, stable URLs, and (where applicable) publication dates on the exhibit index; "
        "archive each URL (for example PDF print-to-file with date) if counsel advises."
    )

    blog_rows = collect_videotext_blog_rows(repo_root)
    extra_rows = load_extra_public_writing(script_dir)
    all_rows = blog_rows + extra_rows

    if all_rows:
        add_heading(doc, "Article index (title, URL, channel)", level=2)
        pub_headers = ["Title", "URL", "Channel / platform"]
        add_table(doc, pub_headers, all_rows, col_widths_inches=[2.0, 2.6, 1.4])
        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(6)
        nr = note.add_run(
            f"Auto-listed VideoText.io blog posts: {len(blog_rows)}. "
            f"Additional lines from extra_public_writing_urls.txt: {len(extra_rows)}."
        )
        nr.italic = True
        nr.font.size = Pt(10)
    else:
        empty = doc.add_paragraph(
            "No blog posts were found under content/blog, and extra_public_writing_urls.txt has no entries. "
            "Add Medium or other URLs to docs/O1A/extra_public_writing_urls.txt (see file comments) and run "
            "the generator again."
        )
        empty.runs[0].italic = True

    # --- Closing ---
    doc.add_paragraph()
    closing = doc.add_paragraph()
    cr = closing.add_run(
        "Prepared as a technical exhibit for immigration counsel; last regenerated 2026-08-08. Attach the "
        "camera-ready ICCS PDF to the petition as the primary scholarly exhibit. The operational snapshot "
        "above is current as of the regeneration date; refresh it from the live dashboard immediately "
        "before filing. Add redacted revenue and subscriber-count figures only as a separate exhibit if "
        "counsel approves — this memorandum otherwise cites operational/quality metrics (throughput, "
        "reliability, ratings, growth) rather than revenue, which is more time-sensitive and commercially "
        "sensitive."
    )
    cr.italic = True
    cr.font.size = Pt(10)

    return doc


def main() -> None:
    out_dir = Path(__file__).resolve().parent
    out_path = out_dir / "RESEARCH_TO_PRODUCT_CRITERION3_TECHNICAL_SUMMARY.docx"
    doc = build_document(script_dir=out_dir)
    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
